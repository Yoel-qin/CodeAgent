"""Word(.docx) 解析（Phase 1.5a/1.5c，python-docx）。

按段落样式（``Title``/``Heading N``）出 HEADING，余为 PARAGRAPH；表格结构化为
TABLE DocElement（逻辑行 + 横向合并 col_span 经 ``_tc`` 标识识别；纵向合并 vMerge 留后续）。

表格结构化 JSON / HTML / 描述放 ``metadata``，供 doc_chunker 落 ``doc_chunks`` 表格列。
docx 无页码概念（``page_number=None``）。
"""
from __future__ import annotations

import io

from docx import Document

from app.pipeline.parsing.doc_element import DocElement, ParseMeta
from app.pipeline.parsing.table_utils import LogicalRow, logical_to_struct, to_text_dump


def _heading_level(style_name: str) -> int | None:
    name = (style_name or "").lower().strip()
    if name == "title":
        return 1
    if name.startswith("heading"):
        rest = name.replace("heading", "").strip()
        try:
            return max(1, min(6, int(rest)))
        except ValueError:
            return 1
    return None


def _docx_logical_rows(table) -> tuple[list[LogicalRow], int, int]:
    """python-docx 表格 → 逻辑行（横向合并 col_span 经 _tc 标识；vMerge 暂不处理）。"""
    rows: list[LogicalRow] = []
    for tr in table.rows:
        logical: LogicalRow = []
        prev_tc = None
        for cell in tr.cells:
            if cell._tc is prev_tc:
                txt, span = logical[-1]
                logical[-1] = (txt, span + 1)   # 同一 _tc：横向合并，扩 span
            else:
                logical.append((cell.text.strip(), 1))
                prev_tc = cell._tc
        rows.append(logical)
    n_rows = len(table.rows)
    n_cols = max((len(tr.cells) for tr in table.rows), default=0)
    return rows, n_rows, n_cols


def parse_docx(data: bytes, file_path: str) -> tuple[list[DocElement], ParseMeta]:
    document = Document(io.BytesIO(data))
    elements: list[DocElement] = []

    for p in document.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        lvl = _heading_level(p.style.name if p.style else "")
        if lvl is not None:
            elements.append(DocElement(type="HEADING", content=txt, heading_level=lvl))
        else:
            elements.append(DocElement(type="PARAGRAPH", content=txt))

    total_tables = len(document.tables)
    for tbl in document.tables:
        logical, n_rows, n_cols = _docx_logical_rows(tbl)
        table_data, html, desc = logical_to_struct(logical, n_rows, n_cols)
        elements.append(DocElement(
            type="TABLE",
            content=f"{desc}\n{to_text_dump(logical)}",
            metadata={"table_data": table_data, "table_html": html,
                      "table_description": desc, "n_rows": n_rows, "n_cols": n_cols},
        ))

    # 图片（Phase 1.5b，python-docx 内嵌图片：遍历 image 关系取 blob）
    total_images = 0
    for rel in document.part.rels.values():
        try:
            if not rel.reltype.endswith("/image"):
                continue
            part = rel.target_part
            raw = part.blob
            if not raw:
                continue
            content_type = getattr(part, "content_type", "") or ""
            ext = content_type.split("/")[-1] if "/" in content_type else "png"
            elements.append(DocElement(
                type="IMAGE", content="",
                metadata={"image_bytes": raw, "ext": ext, "width": None, "height": None},
            ))
            total_images += 1
        except Exception:
            continue

    return elements, ParseMeta(file_format="docx", parse_engine="python-docx",
                               total_tables=total_tables, total_images=total_images)
