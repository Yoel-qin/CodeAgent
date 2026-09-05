import fitz  # pymupdf
from docx import Document as DocxDocument

from app.pipeline.parsing import parse_doc


def test_markdown_headings_and_paragraph():
    md = "# 标题一\n\n段落甲。\n\n## 小节\n\n段落乙。\n".encode()
    elems, meta = parse_doc(md, ".md", "guide.md")
    assert meta.file_format == "markdown" and meta.parse_status == "COMPLETED"
    types = [e.type for e in elems]
    assert "HEADING" in types and "PARAGRAPH" in types
    para = next(e for e in elems if e.type == "PARAGRAPH" and "段落甲" in e.content)
    assert para.heading_path == ["标题一"]


def test_txt_plain():
    elems, meta = parse_doc(b"hello\nworld", ".txt", "a.txt")
    assert meta.parse_engine == "text" and elems


def test_pdf_generated_in_test(tmp_path):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "PDF section text line")
    data = doc.tobytes()
    elems, meta = parse_doc(data, ".pdf", "a.pdf")
    assert meta.file_format == "pdf" and meta.total_pages == 1
    assert any("PDF section text" in e.content for e in elems if e.type == "PARAGRAPH")


def test_docx_generated_in_test(tmp_path):
    d = DocxDocument()
    d.add_heading("Doc Title", level=1)
    d.add_paragraph("Body text here.")
    p = tmp_path / "a.docx"
    d.save(str(p))
    elems, meta = parse_doc(p.read_bytes(), ".docx", "a.docx")
    assert meta.file_format == "docx"
    assert any("Doc Title" in (e.content or "") for e in elems if e.type == "HEADING")


def test_xlsx_generated_in_test(tmp_path):
    """KEEP xlsx：多 sheet → 每 sheet 一个 HEADING + 一个结构化 TABLE；空 sheet 跳过。"""
    from openpyxl import Workbook

    from app.pipeline.parsing import parse_doc

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "用户统计"
    ws1.append(["姓名", "部门", "登录次数"])
    ws1.append(["张三", "研发部", 12])
    ws1.append(["李四", "运维部", 5])
    wb.create_sheet("空表")
    p = tmp_path / "stats.xlsx"
    wb.save(p)

    elements, meta = parse_doc(p.read_bytes(), ".xlsx", str(p))
    assert meta.file_format == "xlsx" and meta.parse_status == "COMPLETED"
    assert meta.total_tables == 1  # 只有非空 sheet 计入
    kinds = [(e.type, e.content) for e in elements]
    assert kinds[0] == ("HEADING", "用户统计")
    tables = [e for e in elements if e.type == "TABLE"]
    assert len(tables) == 1
    td = tables[0].metadata["table_data"]
    assert td["headers"] == ["姓名", "部门", "登录次数"]
    assert td["rows"][0][0] == "张三" and td["n_rows"] == 3 and td["n_cols"] == 3
    assert "张三" in tables[0].content
