"""Phase 1.5c 表格单测：python-docx 表格结构化（含横向合并 col_span）/ table_utils /
doc_chunker 表格切片 + 大表分片 / _to_orm 透传表格列。无基础设施依赖。"""
from __future__ import annotations

from docx import Document

from app.pipeline.chunking.doc_chunker import chunk_doc_elements
from app.pipeline.ingest_doc import _to_orm
from app.pipeline.parsing.doc_element import DocElement
from app.pipeline.parsing.table_utils import grid_to_logical, logical_to_struct
from app.pipeline.parsing.word_parser import parse_docx

# ---------- python-docx 表格结构化 ----------

def test_docx_table_structured_with_colspan(tmp_path):
    p = tmp_path / "t.docx"
    d = Document()
    d.add_heading("Doc", level=1)
    t = d.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            t.cell(r, c).text = f"r{r}c{c}"
    t.cell(0, 0).merge(t.cell(0, 1))   # 横向合并 → col_span=2
    d.save(str(p))

    elements, meta = parse_docx(p.read_bytes(), str(p))
    tables = [e for e in elements if e.type == "TABLE"]
    assert len(tables) == 1 and meta.total_tables == 1
    td = tables[0].metadata["table_data"]
    assert meta.total_tables == 1
    assert td["n_rows"] == 3 and td["n_cols"] == 3
    assert td["col_spans"][0] == [2, 1]          # 首行合并单元格 span=2
    assert len(td["headers"]) == 2               # 合并后首行 2 个逻辑单元格
    assert "r0c0" in td["headers"][0] and "r0c1" in td["headers"][0]
    assert "<table>" in tables[0].metadata["table_html"]
    assert "colspan" in tables[0].metadata["table_html"]


# ---------- table_utils（PDF 网格路径） ----------

def test_grid_to_logical_and_struct():
    grid = [["A", "B", None], ["1", None, "3"]]
    logical = grid_to_logical(grid)
    assert logical == [[("A", 1), ("B", 1), ("", 1)], [("1", 1), ("", 1), ("3", 1)]]
    data, html, desc = logical_to_struct(logical, n_rows=2, n_cols=3)
    assert data["headers"] == ["A", "B", ""] and data["n_rows"] == 2
    assert "<table>" in html and desc.startswith("2 行 × 3 列表格")


# ---------- doc_chunker 表格切片 ----------

def _table_meta(n_rows: int, rows: list[list[str]], n_cols: int = 2) -> dict:
    return {
        "table_data": {"headers": ["A", "B"], "rows": rows, "n_rows": n_rows, "n_cols": n_cols},
        "table_html": "<table></table>", "table_description": "desc",
        "n_rows": n_rows, "n_cols": n_cols,
    }


def test_small_table_one_chunk():
    meta = _table_meta(n_rows=3, rows=[["1", "2"], ["3", "4"]])
    elements = [
        DocElement(type="HEADING", content="Sec", heading_level=2),
        DocElement(type="TABLE", content="desc\nA B\n1 2\n3 4", metadata=meta),
    ]
    specs = chunk_doc_elements(elements, file_path="t.docx", file_hash="h" * 16, commit_hash="C")
    tables = [s for s in specs if s.chunk_content_type == "table"]
    assert len(tables) == 1
    assert tables[0].table_total_rows == 3 and tables[0].table_total_cols == 2
    assert tables[0].table_data["headers"] == ["A", "B"]
    assert tables[0].table_description == "desc"
    assert tables[0].is_table_fragment is None
    assert tables[0].chunk_id.startswith("tbl_")


def test_large_table_fragments():
    body = [[str(i), str(i + 1)] for i in range(0, 60, 2)]   # 30 正文行 → n_rows=31
    meta = _table_meta(n_rows=31, rows=body)
    elements = [DocElement(type="TABLE", content="x", metadata=meta)]
    specs = chunk_doc_elements(elements, file_path="t.docx", file_hash="h" * 16, commit_hash="C")
    frags = [s for s in specs if s.chunk_content_type == "table_fragment"]
    assert len(frags) == 2                       # 30 行 / 20 → 20 + 10
    assert {f.table_fragment_index for f in frags} == {1, 2}
    parents = {f.parent_table_chunk_id for f in frags}
    assert len(parents) == 1                     # 同一父表
    assert all(f.is_table_fragment for f in frags)


def test_to_orm_populates_table_fields():
    meta = _table_meta(n_rows=3, rows=[["1", "2"], ["3", "4"]])
    elements = [DocElement(type="TABLE", content="x", metadata=meta)]
    spec = chunk_doc_elements(elements, file_path="t.docx", file_hash="h" * 16, commit_hash="C")[0]
    orm = _to_orm(spec, file_id=7)
    assert orm.chunk_content_type == "table"
    assert orm.table_total_rows == 3 and orm.table_total_cols == 2
    assert orm.table_data["headers"] == ["A", "B"]
    assert orm.table_html == "<table></table>"
    assert orm.table_description == "desc"
    assert orm.is_table_fragment is None
