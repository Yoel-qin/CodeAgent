"""Phase 1.5a 解析器与路由单测：用 PyMuPDF/python-docx 内存生成夹具，验证
PDF（标题层级+页码+扫描检测）/ docx（标题+段落+表格）/ txt（分段+GBK 解码）/ 路由分发。
无基础设施依赖（不碰 DB/ES/Milvus）。
"""
from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document

from app.pipeline.parsing.pdf_parser import parse_pdf
from app.pipeline.parsing.router import DOC_FORMAT_EXTS, EXT_KIND, parse_doc
from app.pipeline.parsing.txt_parser import decode_text, parse_txt
from app.pipeline.parsing.word_parser import parse_docx

# ---------- 夹具生成 ----------

def _make_pdf(path: Path, *, scanned: bool = False) -> None:
    doc = fitz.open()
    page = doc.new_page()
    if not scanned:
        page.insert_text((72, 72), "Big Heading", fontsize=24)
        page.insert_text((72, 120), "First body paragraph line one.", fontsize=11)
        page.insert_text((72, 138), "Second body paragraph line two.", fontsize=11)
    doc.save(str(path))
    doc.close()


def _make_docx(path: Path) -> None:
    d = Document()
    d.add_heading("Doc Title", level=1)
    d.add_paragraph("Hello body paragraph.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text, t.cell(0, 1).text, t.cell(1, 0).text, t.cell(1, 1).text = "A", "B", "C", "D"
    d.save(str(path))


# ---------- PDF ----------

def test_pdf_heading_and_body(tmp_path):
    p = tmp_path / "spec.pdf"
    _make_pdf(p)
    elements, meta = parse_pdf(p.read_bytes(), str(p))

    headings = [e for e in elements if e.type == "HEADING"]
    paras = [e for e in elements if e.type == "PARAGRAPH"]
    assert len(headings) == 1 and headings[0].content == "Big Heading"
    assert headings[0].heading_level == 1
    assert headings[0].page_number == 1
    assert paras, "应至少一个段落"
    assert "First body" in paras[0].content and "Second body" in paras[0].content
    assert all(e.page_number == 1 for e in paras)
    assert meta.file_format == "pdf" and meta.parse_engine == "pymupdf"
    assert meta.total_pages == 1 and meta.ocr_required is False
    assert meta.parse_status == "COMPLETED"


def test_pdf_scanned_page_flagged(tmp_path):
    p = tmp_path / "scan.pdf"
    _make_pdf(p, scanned=True)
    elements, meta = parse_pdf(p.read_bytes(), str(p))
    assert meta.ocr_required is True
    assert meta.parse_status == "PARTIAL"
    assert elements == []  # 无可抽文本


# ---------- DOCX ----------

def test_docx_heading_paragraph_table(tmp_path):
    p = tmp_path / "spec.docx"
    _make_docx(p)
    elements, meta = parse_docx(p.read_bytes(), str(p))

    types = [e.type for e in elements]
    assert types[0] == "HEADING" and elements[0].heading_level == 1
    assert elements[0].content == "Doc Title"
    assert "PARAGRAPH" in types
    assert "TABLE" in types
    table = next(e for e in elements if e.type == "TABLE")
    assert "A" in table.content and "D" in table.content
    assert meta.file_format == "docx" and meta.total_tables == 1


# ---------- TXT ----------

def test_txt_paragraph_split():
    elements, meta = parse_txt(b"first para\n\nsecond para\n\nthird", "n.txt")
    assert len(elements) == 3 and all(e.type == "PARAGRAPH" for e in elements)
    assert elements[0].content == "first para"
    assert meta.file_format == "txt" and meta.parse_engine == "text"


def test_txt_gb18030_decode():
    raw = "中文段落测试".encode("gb18030")
    assert decode_text(raw) == "中文段落测试"
    elements, _ = parse_txt(raw, "c.txt")
    assert elements and "中文" in elements[0].content


# ---------- 路由 ----------

def test_router_dispatch_unknown_ext_raises():
    with pytest.raises(ValueError):
        parse_doc(b"x", ".xyz", "x.xyz")


def test_router_doc_legacy_failed():
    elements, meta = parse_doc(b"x", ".doc", "legacy.doc")
    assert elements == []
    assert meta.parse_status == "FAILED" and meta.file_format == "doc"


def test_router_markdown_route(tmp_path):
    elements, meta = parse_doc(b"# Title\n\nbody", ".md", "t.md")
    assert meta.file_format == "markdown" and meta.parse_engine == "markdown"
    assert any(e.type == "HEADING" and e.content == "Title" for e in elements)


def test_ext_kind_registry_consistent():
    # code 与 doc 扩展均经唯一真相源；新格式就位
    assert EXT_KIND[".java"] == "code"
    for ext in (".pdf", ".docx", ".txt", ".md"):
        assert EXT_KIND[ext] == "doc"
    assert DOC_FORMAT_EXTS[".pdf"] == "pdf"
    assert DOC_FORMAT_EXTS[".docx"] == "docx"
